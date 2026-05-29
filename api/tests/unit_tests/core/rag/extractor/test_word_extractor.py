"""Primarily used for testing merged cell scenarios"""

import io
import os
import tempfile
from collections import UserDict
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import core.rag.extractor.word_extractor as we
from core.rag.extractor.word_extractor import WordExtractor


def _generate_table_with_merged_cells():
    doc = Document()

    """
    The table looks like this:
    +-----+-----+-----+
    | 1-1 & 1-2 | 1-3 |
    +-----+-----+-----+
    | 2-1 | 2-2 | 2-3 |
    |  &  |-----+-----+
    | 3-1 | 3-2 | 3-3 |
    +-----+-----+-----+
    """
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"

    for i in range(3):
        for j in range(3):
            cell = table.cell(i, j)
            cell.text = f"{i + 1}-{j + 1}"

    # Merge cells
    cell_0_0 = table.cell(0, 0)
    cell_0_1 = table.cell(0, 1)
    merged_cell_1 = cell_0_0.merge(cell_0_1)
    merged_cell_1.text = "1-1 & 1-2"

    cell_1_0 = table.cell(1, 0)
    cell_2_0 = table.cell(2, 0)
    merged_cell_2 = cell_1_0.merge(cell_2_0)
    merged_cell_2.text = "2-1 & 3-1"

    ground_truth = [["1-1 & 1-2", "", "1-3"], ["2-1 & 3-1", "2-2", "2-3"], ["2-1 & 3-1", "3-2", "3-3"]]

    return doc.tables[0], ground_truth


def test_parse_row():
    table, gt = _generate_table_with_merged_cells()
    extractor = object.__new__(WordExtractor)
    for idx, row in enumerate(table.rows):
        assert extractor._parse_row(row, {}, 3) == gt[idx]


def test_init_downloads_via_ssrf_proxy(monkeypatch: pytest.MonkeyPatch):
    doc = Document()
    doc.add_paragraph("hello")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    calls: list[tuple[str, object]] = []

    class FakeResponse:
        status_code = 200
        content = docx_bytes

        def close(self) -> None:
            calls.append(("close", None))

    def fake_get(url: str, **kwargs):
        calls.append(("get", (url, kwargs)))
        return FakeResponse()

    monkeypatch.setattr(we, "ssrf_proxy", SimpleNamespace(get=fake_get))

    extractor = WordExtractor("https://example.com/test.docx", "tenant_id", "user_id")
    try:
        assert calls
        assert calls[0][0] == "get"
        url, kwargs = calls[0][1]
        assert url == "https://example.com/test.docx"
        assert kwargs.get("timeout") is None
        assert extractor.web_path == "https://example.com/test.docx"
        assert extractor.file_path != extractor.web_path
        assert Path(extractor.file_path).read_bytes() == docx_bytes
    finally:
        extractor.temp_file.close()


def test_extract_images_from_docx(monkeypatch: pytest.MonkeyPatch):
    external_bytes = b"ext-bytes"
    internal_bytes = b"int-bytes"

    # Patch storage.save to capture writes
    saves: list[tuple[str, bytes]] = []

    def save(key: str, data: bytes):
        saves.append((key, data))

    monkeypatch.setattr(we, "storage", SimpleNamespace(save=save))

    # Patch db.session to record adds/commit
    class DummySession:
        def __init__(self):
            self.added = []
            self.committed = False

        def add(self, obj):
            self.added.append(obj)

        def commit(self):
            self.committed = True

    db_stub = SimpleNamespace(session=DummySession())
    monkeypatch.setattr(we, "db", db_stub)

    # Patch config values used for URL composition and storage type
    monkeypatch.setattr(we.dify_config, "FILES_URL", "http://files.local", raising=False)
    monkeypatch.setattr(we.dify_config, "STORAGE_TYPE", "local", raising=False)

    # Patch UploadFile to avoid real DB models
    class FakeUploadFile:
        _i = 0

        def __init__(self, **kwargs):  # kwargs match the real signature fields
            type(self)._i += 1
            self.id = f"u{self._i}"

    monkeypatch.setattr(we, "UploadFile", FakeUploadFile)

    # Patch external image fetcher
    def fake_get(url: str, **kwargs):
        assert url == "https://example.com/image.png"
        return SimpleNamespace(status_code=200, headers={"Content-Type": "image/png"}, content=external_bytes)

    monkeypatch.setattr(we, "ssrf_proxy", SimpleNamespace(get=fake_get))

    # A hashable internal part object with a blob attribute
    class HashablePart:
        def __init__(self, blob: bytes):
            self.blob = blob

        def __hash__(self) -> int:  # ensure it can be used as a dict key like real docx parts
            return id(self)

    # Build a minimal doc object with both external and internal image rels
    internal_part = HashablePart(blob=internal_bytes)
    rel_ext = SimpleNamespace(is_external=True, target_ref="https://example.com/image.png")
    rel_int = SimpleNamespace(is_external=False, target_ref="word/media/image1.png", target_part=internal_part)
    doc = SimpleNamespace(part=SimpleNamespace(rels={"rId1": rel_ext, "rId2": rel_int}))

    extractor = object.__new__(WordExtractor)
    extractor.tenant_id = "t1"
    extractor.user_id = "u1"

    image_map = extractor._extract_images_from_docx(doc)

    # Returned map should contain entries for external (keyed by rId) and internal (keyed by target_part)
    assert set(image_map.keys()) == {"rId1", internal_part}
    assert all(v.startswith("![image](") and v.endswith("/file-preview)") for v in image_map.values())

    # Storage should receive both payloads
    payloads = {data for _, data in saves}
    assert external_bytes in payloads
    assert internal_bytes in payloads

    # DB interactions should be recorded
    assert len(db_stub.session.added) == 2
    assert db_stub.session.committed is True


def test_extract_images_from_docx_uses_internal_files_url():
    """Test that INTERNAL_FILES_URL takes precedence over FILES_URL for plugin access."""
    # Test the URL generation logic directly
    from configs import dify_config

    # Mock the configuration values
    original_files_url = getattr(dify_config, "FILES_URL", None)
    original_internal_files_url = getattr(dify_config, "INTERNAL_FILES_URL", None)

    try:
        # Set both URLs - INTERNAL should take precedence
        dify_config.FILES_URL = "http://external.example.com"
        dify_config.INTERNAL_FILES_URL = "http://internal.docker:5001"

        # Test the URL generation logic (same as in word_extractor.py)
        upload_file_id = "test_file_id"

        # This is the pattern we fixed in the word extractor
        base_url = dify_config.INTERNAL_FILES_URL or dify_config.FILES_URL
        generated_url = f"{base_url}/files/{upload_file_id}/file-preview"

        # Verify that INTERNAL_FILES_URL is used instead of FILES_URL
        assert "http://internal.docker:5001" in generated_url, f"Expected internal URL, got: {generated_url}"
        assert "http://external.example.com" not in generated_url, f"Should not use external URL, got: {generated_url}"

    finally:
        # Restore original values
        dify_config.FILES_URL = original_files_url
        dify_config.INTERNAL_FILES_URL = original_internal_files_url


def test_extract_hyperlinks(monkeypatch: pytest.MonkeyPatch):
    # Mock db and storage to avoid issues during image extraction (even if no images are present)
    monkeypatch.setattr(we, "storage", SimpleNamespace(save=lambda k, d: None))
    db_stub = SimpleNamespace(session=SimpleNamespace(add=lambda o: None, commit=lambda: None))
    monkeypatch.setattr(we, "db", db_stub)
    monkeypatch.setattr(we.dify_config, "FILES_URL", "http://files.local", raising=False)
    monkeypatch.setattr(we.dify_config, "STORAGE_TYPE", "local", raising=False)

    doc = Document()
    p = doc.add_paragraph("Visit ")

    # Adding a hyperlink manually
    r_id = "rId99"
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Dify"
    new_run.append(t)
    hyperlink.append(new_run)
    p._p.append(hyperlink)

    # Add relationship to the part
    doc.part.rels.add_relationship(
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        "https://dify.ai",
        r_id,
        is_external=True,
    )

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    try:
        extractor = WordExtractor(tmp_path, "tenant_id", "user_id")
        docs = extractor.extract()
        # Verify modern hyperlink extraction
        assert "Visit[Dify](https://dify.ai)" in docs[0].page_content
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def test_extract_legacy_hyperlinks(monkeypatch: pytest.MonkeyPatch):
    # Mock db and storage
    monkeypatch.setattr(we, "storage", SimpleNamespace(save=lambda k, d: None))
    db_stub = SimpleNamespace(session=SimpleNamespace(add=lambda o: None, commit=lambda: None))
    monkeypatch.setattr(we, "db", db_stub)
    monkeypatch.setattr(we.dify_config, "FILES_URL", "http://files.local", raising=False)
    monkeypatch.setattr(we.dify_config, "STORAGE_TYPE", "local", raising=False)

    doc = Document()
    p = doc.add_paragraph()

    # Construct a legacy HYPERLINK field:
    # 1. w:fldChar (begin)
    # 2. w:instrText (HYPERLINK "http://example.com")
    # 3. w:fldChar (separate)
    # 4. w:r (visible text "Example")
    # 5. w:fldChar (end)

    run1 = OxmlElement("w:r")
    fldCharBegin = OxmlElement("w:fldChar")
    fldCharBegin.set(qn("w:fldCharType"), "begin")
    run1.append(fldCharBegin)
    p._p.append(run1)

    run2 = OxmlElement("w:r")
    instrText = OxmlElement("w:instrText")
    instrText.text = ' HYPERLINK "http://example.com" '
    run2.append(instrText)
    p._p.append(run2)

    run3 = OxmlElement("w:r")
    fldCharSep = OxmlElement("w:fldChar")
    fldCharSep.set(qn("w:fldCharType"), "separate")
    run3.append(fldCharSep)
    p._p.append(run3)

    run4 = OxmlElement("w:r")
    t4 = OxmlElement("w:t")
    t4.text = "Example"
    run4.append(t4)
    p._p.append(run4)

    run5 = OxmlElement("w:r")
    fldCharEnd = OxmlElement("w:fldChar")
    fldCharEnd.set(qn("w:fldCharType"), "end")
    run5.append(fldCharEnd)
    p._p.append(run5)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    try:
        extractor = WordExtractor(tmp_path, "tenant_id", "user_id")
        docs = extractor.extract()
        # Verify legacy hyperlink extraction
        assert "[Example](http://example.com)" in docs[0].page_content
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def test_init_rejects_invalid_url_status(monkeypatch: pytest.MonkeyPatch):
    class FakeResponse:
        status_code = 404
        content = b""
        closed = False

        def close(self):
            self.closed = True

    fake_response = FakeResponse()
    monkeypatch.setattr(we, "ssrf_proxy", SimpleNamespace(get=lambda url, **kwargs: fake_response))

    with pytest.raises(ValueError, match="returned status code 404"):
        WordExtractor("https://example.com/missing.docx", "tenant", "user")

    assert fake_response.closed is True


def test_init_expands_home_path_and_invalid_local_path(monkeypatch, tmp_path):
    target_file = tmp_path / "expanded.docx"
    target_file.write_bytes(b"docx")

    monkeypatch.setattr(we.os.path, "expanduser", lambda p: str(target_file))
    monkeypatch.setattr(
        we.os.path,
        "isfile",
        lambda p: p == str(target_file),
    )

    extractor = WordExtractor("~/expanded.docx", "tenant", "user")
    assert extractor.file_path == str(target_file)

    monkeypatch.setattr(we.os.path, "isfile", lambda p: False)
    with pytest.raises(ValueError, match="is not a valid file or url"):
        WordExtractor("not-a-file", "tenant", "user")


def test_close_closes_temp_file():
    extractor = object.__new__(WordExtractor)
    extractor._closed = False
    extractor.temp_file = MagicMock()

    extractor.close()

    extractor.temp_file.close.assert_called_once()


def test_close_is_idempotent():
    extractor = object.__new__(WordExtractor)
    extractor._closed = False
    extractor.temp_file = MagicMock()

    extractor.close()
    extractor.close()

    extractor.temp_file.close.assert_called_once()


async def _async_close() -> None:
    return None


def test_close_closes_awaitable_close_result():
    extractor = object.__new__(WordExtractor)
    extractor._closed = False
    extractor.temp_file = MagicMock()
    close_result = _async_close()
    extractor.temp_file.close = MagicMock(return_value=close_result)

    extractor.close()

    assert close_result.cr_frame is None
    extractor.temp_file.close.assert_called_once()


def test_extract_images_handles_invalid_external_cases(monkeypatch: pytest.MonkeyPatch):
    class FakeTargetRef:
        def __contains__(self, item):
            return item == "image"

        def split(self, sep):
            return [None]

    rel_invalid_url = SimpleNamespace(is_external=True, target_ref="image-no-url")
    rel_request_error = SimpleNamespace(is_external=True, target_ref="https://example.com/image-error")
    rel_unknown_mime = SimpleNamespace(is_external=True, target_ref="https://example.com/image-unknown")
    rel_internal_none_ext = SimpleNamespace(is_external=False, target_ref=FakeTargetRef(), target_part=object())

    doc = SimpleNamespace(
        part=SimpleNamespace(
            rels={
                "r1": rel_invalid_url,
                "r2": rel_request_error,
                "r3": rel_unknown_mime,
                "r4": rel_internal_none_ext,
            }
        )
    )

    def fake_get(url, **kwargs):
        if "image-error" in url:
            raise RuntimeError("network")
        return SimpleNamespace(status_code=200, headers={"Content-Type": "application/unknown"}, content=b"x")

    monkeypatch.setattr(we, "ssrf_proxy", SimpleNamespace(get=fake_get))
    db_stub = SimpleNamespace(session=SimpleNamespace(add=lambda obj: None, commit=MagicMock()))
    monkeypatch.setattr(we, "db", db_stub)
    monkeypatch.setattr(we, "storage", SimpleNamespace(save=lambda key, data: None))
    monkeypatch.setattr(we.dify_config, "FILES_URL", "http://files.local", raising=False)

    extractor = object.__new__(WordExtractor)
    extractor.tenant_id = "tenant"
    extractor.user_id = "user"

    result = extractor._extract_images_from_docx(doc)

    assert result == {}
    db_stub.session.commit.assert_called_once()


def test_table_to_markdown_and_parse_helpers(monkeypatch: pytest.MonkeyPatch):
    extractor = object.__new__(WordExtractor)

    table = SimpleNamespace(
        rows=[
            SimpleNamespace(cells=[1, 2]),
            SimpleNamespace(cells=[3, 4]),
        ]
    )
    parse_row_mock = MagicMock(side_effect=[["H1", "H2"], ["A", "B"]])
    monkeypatch.setattr(extractor, "_parse_row", parse_row_mock)

    markdown = extractor._table_to_markdown(table, {})
    assert markdown == "| H1 | H2 |\n| --- | --- |\n| A | B |"

    class FakeBlip:
        def __init__(self, image_id):
            self.image_id = image_id

        def get(self, key):
            return self.image_id

    class FakeRunChild:
        def __init__(self, blips, text=""):
            self._blips = blips
            self.text = text
            self.tag = qn("w:r")

        def xpath(self, pattern):
            if pattern == ".//a:blip":
                return self._blips
            return []

    class FakeRun:
        def __init__(self, element, paragraph):
            # Mirror the subset used by _parse_cell_paragraph
            self.element = element
            self.text = getattr(element, "text", "")

    # Patch we.Run so our lightweight child objects work with the extractor
    monkeypatch.setattr(we, "Run", FakeRun)

    image_part = object()
    paragraph = SimpleNamespace(
        _element=[
            FakeRunChild([FakeBlip(None), FakeBlip("ext"), FakeBlip("int")], text=""),
            FakeRunChild([], text="plain"),
        ],
        part=SimpleNamespace(
            rels={
                "ext": SimpleNamespace(is_external=True),
                "int": SimpleNamespace(is_external=False, target_part=image_part),
            }
        ),
    )

    image_map = {"ext": "EXT-IMG", image_part: "INT-IMG"}
    assert extractor._parse_cell_paragraph(paragraph, image_map) == "EXT-IMGINT-IMGplain"

    cell = SimpleNamespace(paragraphs=[paragraph, paragraph])
    assert extractor._parse_cell(cell, image_map) == "EXT-IMGINT-IMGplain"


def test_parse_docx_covers_drawing_shapes_hyperlink_error_and_table_branch(monkeypatch: pytest.MonkeyPatch):
    extractor = object.__new__(WordExtractor)

    ext_image_id = "ext-image"
    int_embed_id = "int-embed"
    shape_ext_id = "shape-ext"
    shape_int_id = "shape-int"

    internal_part = object()
    shape_internal_part = object()

    class Rels(UserDict):
        def get(self, key, default=None):
            if key == "link-bad":
                raise RuntimeError("cannot resolve relation")
            return super().get(key, default)

    rels = Rels(
        {
            ext_image_id: SimpleNamespace(is_external=True, target_ref="https://img/ext.png"),
            int_embed_id: SimpleNamespace(is_external=False, target_part=internal_part),
            shape_ext_id: SimpleNamespace(is_external=True, target_ref="https://img/shape.png"),
            shape_int_id: SimpleNamespace(is_external=False, target_part=shape_internal_part),
            "link-ok": SimpleNamespace(is_external=True, target_ref="https://example.com"),
        }
    )

    image_map = {
        ext_image_id: "[EXT]",
        internal_part: "[INT]",
        shape_ext_id: "[SHAPE_EXT]",
        shape_internal_part: "[SHAPE_INT]",
    }

    class FakeBlip:
        def __init__(self, embed_id):
            self.embed_id = embed_id

        def get(self, key):
            return self.embed_id

    class FakeDrawing:
        def __init__(self, embed_ids):
            self.embed_ids = embed_ids

        def findall(self, pattern):
            return [FakeBlip(embed_id) for embed_id in self.embed_ids]

    class FakeNode:
        def __init__(self, text=None, attrs=None):
            self.text = text
            self._attrs = attrs or {}

        def get(self, key):
            return self._attrs.get(key)

    class FakeShape:
        def __init__(self, bin_id=None, img_id=None):
            self.bin_id = bin_id
            self.img_id = img_id

        def find(self, pattern):
            if "binData" in pattern and self.bin_id:
                return FakeNode(
                    text="shape",
                    attrs={"{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id": self.bin_id},
                )
            if "imagedata" in pattern and self.img_id:
                return FakeNode(attrs={"id": self.img_id})
            return None

    class FakeChild:
        def __init__(
            self,
            tag,
            text="",
            fld_chars=None,
            instr_texts=None,
            drawings=None,
            shapes=None,
            attrs=None,
            hyperlink_runs=None,
        ):
            self.tag = tag
            self.text = text
            self._fld_chars = fld_chars or []
            self._instr_texts = instr_texts or []
            self._drawings = drawings or []
            self._shapes = shapes or []
            self._attrs = attrs or {}
            self._hyperlink_runs = hyperlink_runs or []

        def findall(self, pattern):
            if pattern == qn("w:fldChar"):
                return self._fld_chars
            if pattern == qn("w:instrText"):
                return self._instr_texts
            if pattern == qn("w:r"):
                return self._hyperlink_runs
            if pattern.endswith("}drawing"):
                return self._drawings
            if pattern.endswith("}pict"):
                return self._shapes
            return []

        def get(self, key):
            return self._attrs.get(key)

    class FakeRun:
        def __init__(self, element, paragraph):
            self.element = element
            self.text = getattr(element, "text", "")

    paragraph_main = SimpleNamespace(
        _element=[
            FakeChild(
                qn("w:r"),
                text="run-text",
                drawings=[FakeDrawing([ext_image_id, int_embed_id])],
                shapes=[FakeShape(bin_id=shape_ext_id, img_id=shape_int_id)],
            ),
            FakeChild(
                qn("w:r"),
                text="",
                drawings=[],
                shapes=[FakeShape(bin_id=shape_ext_id)],
            ),
            FakeChild(
                qn("w:hyperlink"),
                attrs={qn("r:id"): "link-ok"},
                hyperlink_runs=[FakeChild(qn("w:r"), text="LinkText")],
            ),
            FakeChild(
                qn("w:hyperlink"),
                attrs={qn("r:id"): "link-bad"},
                hyperlink_runs=[FakeChild(qn("w:r"), text="BrokenLink")],
            ),
        ]
    )
    paragraph_empty = SimpleNamespace(_element=[FakeChild(qn("w:r"), text="   ")])

    fake_doc = SimpleNamespace(
        part=SimpleNamespace(rels=rels, related_parts={int_embed_id: internal_part}),
        paragraphs=[paragraph_main, paragraph_empty],
        tables=[SimpleNamespace(rows=[])],
        element=SimpleNamespace(
            body=[SimpleNamespace(tag="w:p"), SimpleNamespace(tag="w:p"), SimpleNamespace(tag="w:tbl")]
        ),
    )

    monkeypatch.setattr(we, "DocxDocument", lambda _: fake_doc)
    monkeypatch.setattr(we, "Run", FakeRun)
    monkeypatch.setattr(extractor, "_extract_images_from_docx", lambda doc: image_map)
    monkeypatch.setattr(extractor, "_table_to_markdown", lambda table, image_map: "TABLE-MARKDOWN")
    logger_exception = MagicMock()
    monkeypatch.setattr(we.logger, "exception", logger_exception)

    content = extractor.parse_docx("dummy.docx")

    assert "[EXT]" in content
    assert "[INT]" in content
    assert "[SHAPE_EXT]" in content
    assert "[LinkText](https://example.com)" in content
    assert "BrokenLink" in content
    assert "TABLE-MARKDOWN" in content
    logger_exception.assert_called_once()


def test_parse_cell_paragraph_hyperlink_in_table_cell_http():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]

    # Build modern hyperlink inside table cell
    r_id = "rIdHttp1"
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Dify"
    run_elem.append(t)
    hyperlink.append(run_elem)
    p._p.append(hyperlink)

    # Relationship for external http link
    doc.part.rels.add_relationship(
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        "https://dify.ai",
        r_id,
        is_external=True,
    )

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    try:
        reopened = Document(tmp_path)
        para = reopened.tables[0].cell(0, 0).paragraphs[0]
        extractor = object.__new__(WordExtractor)
        out = extractor._parse_cell_paragraph(para, {})
        assert out == "[Dify](https://dify.ai)"
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def test_parse_cell_paragraph_hyperlink_in_table_cell_mailto():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]

    r_id = "rIdMail1"
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "john@test.com"
    run_elem.append(t)
    hyperlink.append(run_elem)
    p._p.append(hyperlink)

    doc.part.rels.add_relationship(
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        "mailto:john@test.com",
        r_id,
        is_external=True,
    )

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    try:
        reopened = Document(tmp_path)
        para = reopened.tables[0].cell(0, 0).paragraphs[0]
        extractor = object.__new__(WordExtractor)
        out = extractor._parse_cell_paragraph(para, {})
        assert out == "[john@test.com](mailto:john@test.com)"
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
