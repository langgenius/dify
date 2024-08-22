"""Abstract interface for document loader implementations."""
import os
import tempfile
from urllib.parse import urlparse

import requests
from docx import Document as DocxDocument

from core.rag.extractor.extractor_base import BaseExtractor
from core.rag.models.document import Document


class WordExtractorTest(BaseExtractor):
    """Load docx files.


    Args:
        file_path: Path to the file to load.
    """

    def __init__(self, file_path: str):
        """Initialize with file path."""
        self.file_path = file_path
        if "~" in self.file_path:
            self.file_path = os.path.expanduser(self.file_path)

        # If the file is a web path, download it to a temporary file, and use that
        if not os.path.isfile(self.file_path) and self._is_valid_url(self.file_path):
            r = requests.get(self.file_path)

            if r.status_code != 200:
                raise ValueError(
                    f"Check the url of your file; returned status code {r.status_code}"
                )

            self.web_path = self.file_path
            self.temp_file = tempfile.NamedTemporaryFile()
            self.temp_file.write(r.content)
            self.file_path = self.temp_file.name
        elif not os.path.isfile(self.file_path):
            raise ValueError(f"File path {self.file_path} is not a valid file or url")

    def __del__(self) -> None:
        if hasattr(self, "temp_file"):
            self.temp_file.close()

    def extract(self) -> list[Document]:
        """Load given path as single page."""
        from docx import Document as docx_Document

        document = docx_Document(self.file_path)
        doc_texts = [paragraph.text for paragraph in document.paragraphs]
        content = '\n'.join(doc_texts)

        return [Document(
            page_content=content,
            metadata={"source": self.file_path},
        )]

    @staticmethod
    def _is_valid_url(url: str) -> bool:
        """Check if the url is valid."""
        parsed = urlparse(url)
        return bool(parsed.netloc) and bool(parsed.scheme)

    def _extract_images_from_docx(self, doc, image_folder):
        image_count = 0
        image_paths = []

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                image_ext = rel.target_ref.split('.')[-1]
                image_name = f"image{image_count}.{image_ext}"
                image_path = os.path.join(image_folder, image_name)
                with open(image_path, "wb") as img_file:
                    img_file.write(rel.target_part.blob)
                image_paths.append(f"![](/api/system/img/{image_name})")

        return image_paths

    def _table_to_html(self, table):
        html = "<table border='1'>"
        for row in table.rows:
            html += "<tr>"
            for cell in row.cells:
                html += f"<td>{cell.text}</td>"
            html += "</tr>"
        html += "</table>"
        return html

    def parse_docx(self, docx_path, image_folder):
        doc = DocxDocument(docx_path)
        os.makedirs(image_folder, exist_ok=True)

        content = []

        image_index = 0
        image_paths = self._extract_images_from_docx(doc, image_folder)

        for element in doc.element.body:
            if element.tag.endswith('p'):  # paragraph
                paragraph = element.text.strip()
                if paragraph:
                    content.append(paragraph)
            elif element.tag.endswith('tbl'):  # table
                table = doc.tables[image_index]
                content.append(self._table_to_html(table))
                image_index += 1

        # 替换图片占位符
        content_with_images = []
        for item in content:
            if '!' in item and '[]' in item:
                item = image_paths.pop(0)
            content_with_images.append(item)

        return content_with_images
